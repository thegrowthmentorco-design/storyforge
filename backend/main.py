import io
import json
import logging
import os
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Annotated

from dotenv import load_dotenv

# Load backend/.env BEFORE importing extract (which reads ANTHROPIC_API_KEY)
load_dotenv(Path(__file__).parent / ".env")

import anthropic
from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from pypdf import PdfReader
from docx import Document
from sqlmodel import Session

from auth.deps import (
    CurrentUser,
    current_user,
    enforce_token_rate_limit,
    enforce_token_scope,
)
from db.session import get_session, init_db
from models import ExtractionRecord
from routers import api_tokens as api_tokens_router
from routers import billing as billing_router
from routers import comments as comments_router
from routers import chat as chat_router
from routers import extractions as extractions_router
from routers import integrations as integrations_router
from routers import me as me_router
from routers import projects as projects_router
from routers import share as share_router
from services.byok import resolve_user_byok
from services.extractions import (
    call_claude,
    extraction_to_record,
    mint_extraction_id,
    persist_extraction,
    record_usage,
    save_upload,
)
from services.few_shot import resolve_enabled_examples
from services.ingest import combine_raw_texts, ingest_file
from services.limits import enforce_limits
from services.prompts import resolve_prompt_suffix
from services.obs import install_json_logging, install_request_id, install_sentry
from services.onboarding import welcome_check
from services.streaming import stream_extraction

MAX_BYTES = 10 * 1024 * 1024  # 10 MB
SUPPORTED_EXT = {
    # Text-extractable formats — pypdf / python-docx / utf-8 decode
    ".pdf", ".docx", ".txt", ".md", ".markdown", ".rst",
    # M7.4 — image inputs go through Claude vision (services/vision.py).
    # MIME mapping handled in vision.mime_for_ext.
    ".png", ".jpg", ".jpeg", ".gif", ".webp",
}

# Install JSON logging + Sentry BEFORE FastAPI app construction so any
# import-time log lines from routers/services land in the structured pipeline
# and any startup error is reported to Sentry.
install_json_logging()
install_sentry()
log = logging.getLogger("storyforge")

@asynccontextmanager
async def lifespan(_app: FastAPI):
    init_db()
    yield


# M6.7.d — interactive OpenAPI docs at /api-docs (Swagger UI). The raw
# spec stays at /openapi.json so client codegen tools find it via the
# FastAPI convention. Redoc is disabled — Swagger UI's "Try it out" + the
# Authorize button are what we want for users testing API tokens. The
# docs page is unauthenticated (the spec describes the public API; secrets
# only flow when the user pastes a Bearer token via Authorize).
app = FastAPI(
    title="StoryForge API",
    version="0.3.0",
    description=(
        "Programmatic access to StoryForge extractions, projects, and gap state. "
        "Authenticate with an API token (Settings → API tokens → Create) by clicking "
        "**Authorize** above and pasting `Bearer sk_live_…`. "
        "Read-only tokens (M6.7.b) can call GET endpoints only. Per-token rate limit "
        "is 60 requests/min by default (M6.7.c)."
    ),
    docs_url="/api-docs",
    redoc_url=None,
    lifespan=lifespan,
)

# Request id + access logging. Installed BEFORE CORS so the id is bound for
# the preflight OPTIONS too — useful when debugging cross-origin failures.
install_request_id(app)

# CORS — local dev origins always allowed; prod adds anything in CORS_ORIGINS
# (comma-separated). On Render's single-container deploy the SPA is served
# from the same origin as the API so CORS isn't even triggered at runtime,
# but keeping it tight protects us if someone points a third-party frontend
# at the backend.
_default_cors = ["http://localhost:5173", "http://127.0.0.1:5173"]
_extra_cors = [o.strip() for o in (os.environ.get("CORS_ORIGINS") or "").split(",") if o.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=_default_cors + _extra_cors,
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Router-level auth: every route under /api/extractions, /api/projects, /api/me
# requires a verified Clerk session. /api/health stays public for infra probes.
# `welcome_check` piggybacks on the same dependency chain — fires the welcome
# email exactly once per user, on whichever protected request lands first.
# M6.7.b: `enforce_token_scope` rejects non-safe HTTP methods when the
# caller is using a read-only API token; runs after current_user (which
# stamps the token_scope on the CurrentUser snapshot).
_protected_deps = [
    Depends(current_user),
    Depends(enforce_token_scope),
    Depends(enforce_token_rate_limit),   # M6.7.c — no-op for Clerk sessions
    Depends(welcome_check),
]
app.include_router(extractions_router.router, dependencies=_protected_deps)
# M14.4 — per-extraction chat thread. Same auth posture as extractions.
app.include_router(chat_router.router, dependencies=_protected_deps)
app.include_router(projects_router.router, dependencies=_protected_deps)
app.include_router(me_router.router, dependencies=_protected_deps)
app.include_router(comments_router.router, dependencies=_protected_deps)
app.include_router(integrations_router.router, dependencies=_protected_deps)
# M6.2.d — Jira OAuth callback is unauthenticated at the route layer
# (Atlassian's browser redirect arrives without an Authorization header).
# CSRF state validates the user inside the handler.
app.include_router(integrations_router.unauth_router)
app.include_router(api_tokens_router.router, dependencies=_protected_deps)
# M4.6: share has split posture — owner endpoints are auth+ownership, the
# public read uses token only. Mounted as two separate routers in share.py.
app.include_router(share_router.owner_router, dependencies=_protected_deps)
app.include_router(share_router.public_router)
# Billing router has its own auth posture: /api/me/* routes require Clerk auth
# AND the welcome_check, but /api/webhooks/lemonsqueezy is unauthed (signed
# by LSQ via HMAC). Wired per-route inside the router rather than at this layer.
app.include_router(billing_router.router)


def _parse_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _parse_docx(data: bytes) -> str:
    """Extract paragraph + table cell text from a .docx file."""
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    # Paragraphs in document body order. python-docx doesn't iterate tables and
    # paragraphs in mixed order out of the box, so do paragraphs first, then
    # tables — good enough for BRDs which are mostly prose with occasional tables.
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts).strip()


def _parse_file(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return _parse_pdf(data)
    if lower.endswith(".docx"):
        return _parse_docx(data)
    # Default: treat as text. errors='replace' so a stray binary byte doesn't 500.
    return data.decode("utf-8", errors="replace")


@app.get("/api/health")
def health():
    return {"ok": True, "live": bool(os.environ.get("ANTHROPIC_API_KEY"))}


@app.post("/api/test-key", dependencies=[Depends(enforce_token_scope), Depends(enforce_token_rate_limit), Depends(welcome_check)])
def test_key(
    _user: Annotated[CurrentUser, Depends(current_user)],
    x_anthropic_key: str | None = Header(default=None, alias="X-Anthropic-Key"),
):
    """Validate an Anthropic API key by making a tiny authenticated call.

    Used by the Settings page "Test connection" button. Falls back to the
    server's env key when no header is provided. Cost: one models.list call,
    no token usage.
    """
    key = x_anthropic_key or os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        raise HTTPException(status_code=400, detail="No API key provided.")
    try:
        client = anthropic.Anthropic(api_key=key)
        models = client.models.list()
        count = len(getattr(models, "data", []) or [])
        return {"ok": True, "models_visible": count, "source": "byok" if x_anthropic_key else "env"}
    except anthropic.AuthenticationError:
        raise HTTPException(status_code=401, detail="Invalid API key.")
    except anthropic.PermissionDeniedError:
        raise HTTPException(status_code=403, detail="Key lacks required permissions.")
    except anthropic.APIConnectionError:
        raise HTTPException(status_code=503, detail="Could not reach Anthropic.")
    except anthropic.APIStatusError as e:
        raise HTTPException(status_code=502, detail=f"Anthropic error ({e.status_code}): {e.message}")
    except Exception as e:
        log.exception("test_key failed")
        raise HTTPException(status_code=500, detail=f"Test failed: {e}")


@app.post("/api/extract", response_model=ExtractionRecord, dependencies=[Depends(enforce_token_scope), Depends(enforce_token_rate_limit), Depends(welcome_check)])
async def extract(
    session: Annotated[Session, Depends(get_session)],
    user: Annotated[CurrentUser, Depends(current_user)],
    # M7.5: list[UploadFile] — FastAPI parses repeated `file` form fields
    # into a list. Single-file uploads (one `file` field) → 1-element list,
    # so the existing client doesn't break. Multi-file uploads append the
    # `file` field N times.
    file: list[UploadFile] | None = File(default=None),
    text: str | None = Form(default=None),
    filename: str | None = Form(default=None),
    project_id: str | None = Form(default=None),
    # M14.1 — lens dispatcher. Default 'dossier' (the new narrated 4-act
    # dossier); 'stories' for the legacy user-stories extraction (back-
    # compat with existing API consumers + sub-paths that haven't been
    # ported). Unknown values normalize to the default.
    lens: str | None = Form(default=None),
    x_anthropic_key: str | None = Header(default=None, alias="X-Anthropic-Key"),
    x_storyforge_model: str | None = Header(default=None, alias="X-Storyforge-Model"),
) -> ExtractionRecord:
    files = [f for f in (file or []) if f and f.filename]
    if not files and not text:
        raise HTTPException(status_code=400, detail="Provide either a file or text.")

    # Validate project ownership BEFORE any Claude call — no point burning
    # tokens (especially the per-file vision/OCR pre-passes) if the request
    # is going to fail validation anyway.
    if project_id:
        from db.models import Project as ProjectModel
        from services.scope import in_scope
        proj = session.get(ProjectModel, project_id)
        if not in_scope(proj, user):
            raise HTTPException(status_code=400, detail="Unknown project_id")

    # M3.4.5: pull stored BYOK + model from UserSettings if the request didn't
    # supply them via header. Header still wins (lets users test a new key).
    effective_key, stored_model = resolve_user_byok(session, user.user_id, x_anthropic_key)
    effective_model = x_storyforge_model or stored_model

    # ----- ingest each file (M7.5) -----------------------------------------
    # Loop calls services/ingest.ingest_file, which handles
    # parse + image vision + OCR fallback + usage_log per pre-pass.
    # M7.5.b: persist every uploaded file (not just single-file) so the
    # studio can offer per-doc downloads. We collect (name, bytes) tuples
    # here and write them to R2 / disk after Claude succeeds (so failed
    # extractions don't strand uploads).
    pending_uploads: list[tuple[str, bytes]] = []
    if files:
        per_doc: list[tuple[str, str]] = []
        for f in files:
            text_chunk, name, data, _modality = await ingest_file(
                f,
                session=session, user=user,
                effective_key=effective_key, effective_model=effective_model,
                max_bytes=MAX_BYTES, supported_ext=SUPPORTED_EXT,
                parse_text=_parse_file,
            )
            per_doc.append((name, text_chunk))
            pending_uploads.append((name, data))
        raw_text, source_name = combine_raw_texts(per_doc)
    else:
        raw_text = text or ""
        source_name = filename or "pasted_text.txt"

    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="No readable text in the input.")

    # M3.5.4: gate BEFORE the Claude call — never burn tokens on a doomed
    # request. Raises HTTPException with paywall payload on any tier breach.
    enforce_limits(session, user, raw_text=raw_text, model=effective_model)

    # Anthropic errors are translated to HTTPExceptions inside `call_claude`,
    # so we let them propagate uncaught.
    # M14.1 — normalise lens choice. Unknown values default to dossier.
    from services.lenses import normalize as normalize_lens
    effective_lens = normalize_lens(lens)

    suffix = resolve_prompt_suffix(session, user.user_id, user.org_id)  # M7.1
    examples = resolve_enabled_examples(session, user.user_id, user.org_id)  # M7.2
    result, model_used, usage = call_claude(
        filename=source_name,
        raw_text=raw_text,
        api_key=effective_key,
        model=effective_model,
        prompt_suffix=suffix,
        few_shot_examples=examples,
        lens=effective_lens,
    )

    # Mint the id up front so the upload path can reference it before the
    # row exists. If a disk write fails, we 500 without persisting — no
    # orphaned row pointing at a missing file.
    extraction_id = mint_extraction_id()
    source_path: str | None = None        # legacy single-file column
    source_paths: list[str] = []          # M7.5.b — every uploaded file
    if pending_uploads:
        try:
            for name, data in pending_uploads:
                source_paths.append(save_upload(extraction_id, name, data))
        except OSError as e:
            log.exception("upload save failed")
            raise HTTPException(status_code=500, detail=f"Could not store uploaded file: {e}")
        # Mirror the first path into the legacy column for back-compat with
        # any reader that still keys off `source_file_path`.
        source_path = source_paths[0]

    # M14.1 — branch the persistence by lens. Stories rows go through
    # persist_extraction (existing path); dossier rows go through
    # persist_dossier_extraction which writes lens_payload + folds the
    # user_stories list back into the legacy stories column.
    if effective_lens == "dossier":
        from services.extractions import persist_dossier_extraction
        row = persist_dossier_extraction(
            session,
            filename=source_name,
            raw_text=raw_text,
            dossier=result,
            model_used=model_used,
            live=usage is not None,  # mock dossier returns usage=None
            user_id=user.user_id,
            org_id=user.org_id,
            project_id=project_id or None,
            extraction_id=extraction_id,
            source_file_path=source_path,
            source_file_paths=source_paths,
        )
        live_flag = usage is not None
    else:
        row = persist_extraction(
            session,
            result=result,
            model_used=model_used,
            user_id=user.user_id,
            org_id=user.org_id,
            project_id=project_id or None,
            extraction_id=extraction_id,
            source_file_path=source_path,
            source_file_paths=source_paths,
        )
        live_flag = result.live

    record_usage(
        session,
        user_id=user.user_id,
        org_id=user.org_id,
        extraction_id=row.id,
        action="extract",
        model=model_used,
        live=live_flag,
        usage=usage,
    )
    return extraction_to_record(row)


# ---- Streaming extraction (M5.3) ------------------------------------------
#
# Same request shape as /api/extract; same pre-flight (ownership / paywall /
# input parse). Differences:
#   * The Claude call streams instead of blocking. The route emits Server-Sent
#     Events as the model writes — `usage` events with cumulative input/output
#     token counts so the UI can show a real progress bar; one final `complete`
#     event with the full ExtractionRecord.
#   * Errors during the Claude call land as `error` SSE events (not HTTP
#     errors), because the response status was already 200 by the time
#     streaming started. Pre-flight errors still come back as normal 4xx.
#
# We persist + record_usage *inside* the generator, after streaming completes
# and before the final SSE frame, so the frontend's `complete` payload carries
# the canonical persisted record (with id + created_at) and the usage row is
# in the DB before /api/me/plan returns its next answer.


def _sse(event: str, data: dict | str) -> bytes:
    """Format one SSE frame. Two newlines terminate; we always include both
    `event:` and `data:` so clients can dispatch on event type."""
    payload = data if isinstance(data, str) else json.dumps(data, default=str)
    return f"event: {event}\ndata: {payload}\n\n".encode("utf-8")


@app.post("/api/extract/stream", dependencies=[Depends(enforce_token_scope), Depends(enforce_token_rate_limit), Depends(welcome_check)])
async def extract_stream(
    session: Annotated[Session, Depends(get_session)],
    user: Annotated[CurrentUser, Depends(current_user)],
    file: list[UploadFile] | None = File(default=None),  # M7.5 — see /api/extract
    text: str | None = Form(default=None),
    filename: str | None = Form(default=None),
    project_id: str | None = Form(default=None),
    # M14.1.c — lens dispatcher; default 'dossier' for new uploads.
    lens: str | None = Form(default=None),
    x_anthropic_key: str | None = Header(default=None, alias="X-Anthropic-Key"),
    x_storyforge_model: str | None = Header(default=None, alias="X-Storyforge-Model"),
):
    # IMPORTANT: the request-scoped `session` above is for pre-flight only.
    # FastAPI tears down `Depends(get_session)` right after this handler
    # returns the StreamingResponse object — *before* Starlette iterates the
    # generator to send the body. So inside `event_gen` we open a *fresh*
    # session via `Session(engine)`. This was the M5.3 regression that made
    # extraction silently hang: persist_extraction would run against the
    # closed session, raise inside the generator, and never emit `complete`.
    # ----- pre-flight: identical guards to /api/extract --------------------
    files = [f for f in (file or []) if f and f.filename]
    if not files and not text:
        raise HTTPException(status_code=400, detail="Provide either a file or text.")

    if project_id:
        from db.models import Project as ProjectModel
        from services.scope import in_scope
        proj = session.get(ProjectModel, project_id)
        if not in_scope(proj, user):
            raise HTTPException(status_code=400, detail="Unknown project_id")

    effective_key, stored_model = resolve_user_byok(session, user.user_id, x_anthropic_key)
    effective_model = x_storyforge_model or stored_model

    # ----- ingest each file (M7.5) -----------------------------------------
    # M7.5.b: persist every file (see /api/extract for the rationale).
    pending_uploads: list[tuple[str, bytes]] = []
    if files:
        per_doc: list[tuple[str, str]] = []
        for f in files:
            text_chunk, name, data, _modality = await ingest_file(
                f,
                session=session, user=user,
                effective_key=effective_key, effective_model=effective_model,
                max_bytes=MAX_BYTES, supported_ext=SUPPORTED_EXT,
                parse_text=_parse_file,
            )
            per_doc.append((name, text_chunk))
            pending_uploads.append((name, data))
        raw_text, source_name = combine_raw_texts(per_doc)
    else:
        raw_text = text or ""
        source_name = filename or "pasted_text.txt"

    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="No readable text in the input.")

    # Plan limits — still raise as HTTP errors so the existing paywall modal
    # path catches them. These fire before the SSE stream opens.
    enforce_limits(session, user, raw_text=raw_text, model=effective_model)

    # M7.1 — resolve the user's saved prompt suffix during pre-flight (while
    # the request-scoped session is still alive). Snapshot into a local so
    # the generator's fresh session doesn't have to refetch.
    effective_suffix = resolve_prompt_suffix(session, user.user_id, user.org_id)
    # M7.2 — same lifecycle reasoning for few-shot examples. Snapshot a
    # plain list of (id, name, input_text, expected_payload) so the
    # generator doesn't reach into the closing session.
    effective_examples = resolve_enabled_examples(session, user.user_id, user.org_id)

    # Mint id up front so the start event can carry it (the frontend uses it
    # to wire the in-flight extraction to the persisted row on `complete`).
    extraction_id = mint_extraction_id()
    source_path: str | None = None
    source_paths: list[str] = []
    if pending_uploads:
        try:
            for name, data in pending_uploads:
                source_paths.append(save_upload(extraction_id, name, data))
        except OSError as e:
            log.exception("upload save failed")
            raise HTTPException(status_code=500, detail=f"Could not store uploaded file: {e}")
        source_path = source_paths[0]

    # Snapshot the values we need inside the generator so we don't hold a
    # closure over `session` (which FastAPI is about to close).
    _user_id = user.user_id
    _org_id = user.org_id
    _project_id = project_id or None

    # M14.1.c — normalise lens choice; pick the right streamer + persister.
    from services.lenses import normalize as normalize_lens
    effective_lens = normalize_lens(lens)

    # ----- the SSE generator ------------------------------------------------
    def event_gen():
        from db.session import engine as _engine  # local import to avoid main.py import-time cycle
        from sqlmodel import Session as _Session

        yield _sse("start", {"id": extraction_id, "filename": source_name})
        try:
            # M14.1.c — branch the streamer by lens. Both yield identical SSE
            # event shapes (start / usage / complete / error) so the client
            # doesn't need to know which lens is running.
            if effective_lens == "dossier":
                from services.streaming_dossier import stream_dossier_extraction
                stream_iter = stream_dossier_extraction(
                    filename=source_name,
                    raw_text=raw_text,
                    api_key=effective_key,
                    model=effective_model,
                    prompt_suffix=effective_suffix,
                )
            else:
                stream_iter = stream_extraction(
                    filename=source_name,
                    raw_text=raw_text,
                    api_key=effective_key,
                    model=effective_model,
                    prompt_suffix=effective_suffix,
                    few_shot_examples=effective_examples,
                )

            for ev in stream_iter:
                etype = ev["type"]
                if etype == "usage":
                    yield _sse("usage", {"input": ev["input"], "output": ev["output"], "max": ev["max"]})
                elif etype == "section_ready":
                    # M14.14 — progressive section reveal. Frontend uses these
                    # to mount sections of the dossier as they finish streaming
                    # rather than waiting for the full payload.
                    yield _sse("section_ready", {"key": ev["key"], "value": ev["value"]})
                elif etype == "error":
                    # Stream-time error — surface to client + stop. No persistence.
                    yield _sse("error", {"status": ev["status"], "detail": ev["detail"]})
                    return
                elif etype == "complete":
                    # Open a *fresh* session for persistence — the request-
                    # scoped one is closed by now (see note at the top of the
                    # handler).
                    with _Session(_engine) as s:
                        if effective_lens == "dossier":
                            from services.extractions import persist_dossier_extraction
                            row = persist_dossier_extraction(
                                s,
                                filename=source_name,
                                raw_text=raw_text,
                                dossier=ev["result"],
                                model_used=ev["model_used"],
                                live=ev["usage"] is not None,
                                user_id=_user_id,
                                org_id=_org_id,
                                project_id=_project_id,
                                extraction_id=extraction_id,
                                source_file_path=source_path,
                                source_file_paths=source_paths,
                            )
                            live_flag = ev["usage"] is not None
                        else:
                            row = persist_extraction(
                                s,
                                result=ev["result"],
                                model_used=ev["model_used"],
                                user_id=_user_id,
                                org_id=_org_id,
                                project_id=_project_id,
                                extraction_id=extraction_id,
                                source_file_path=source_path,
                                source_file_paths=source_paths,
                            )
                            live_flag = ev["result"].live
                        record_usage(
                            s,
                            user_id=_user_id,
                            org_id=_org_id,
                            extraction_id=row.id,
                            action="extract",
                            model=ev["model_used"],
                            live=live_flag,
                            usage=ev["usage"],
                        )
                        record = extraction_to_record(row)
                    yield _sse("complete", record.model_dump(mode="json"))
        except Exception as e:
            # Last-resort safety net — anything inside the generator that
            # isn't already converted to an `error` event lands here.
            log.exception("extract_stream generator crashed")
            yield _sse("error", {"status": 500, "detail": f"Extraction failed: {e}"})

    return StreamingResponse(
        event_gen(),
        media_type="text/event-stream",
        # Disable proxy buffering so each frame ships as soon as we yield it.
        # Render's load balancer respects this; nginx/Cloudflare also honour it.
        headers={"X-Accel-Buffering": "no", "Cache-Control": "no-cache"},
    )


# M6.7.d — custom OpenAPI generator that declares a `BearerAuth` security
# scheme so Swagger UI's Authorize button works against our API-token /
# Clerk-JWT auth posture. We don't decorate every route with `Security(...)`
# because the protection is dependency-driven (`_protected_deps`); instead
# we attach the security requirement to every path under /api/ except a
# small public-by-design allowlist (health, share read, billing webhook).
def _custom_openapi():
    if app.openapi_schema:
        return app.openapi_schema
    from fastapi.openapi.utils import get_openapi
    schema = get_openapi(
        title=app.title, version=app.version, description=app.description, routes=app.routes,
    )
    schema.setdefault("components", {}).setdefault("securitySchemes", {})["BearerAuth"] = {
        "type": "http",
        "scheme": "bearer",
        "description": (
            "Paste `sk_live_…` (StoryForge API token from Settings → API tokens) or a "
            "Clerk session JWT. Read-only tokens can only call GET endpoints; per-token "
            "rate limit is 60 req/min."
        ),
    }
    # Public routes — don't decorate them with the security requirement so
    # Swagger doesn't show the (deceptive) lock icon next to them.
    PUBLIC_PATHS = {"/api/health", "/api/share/{token}", "/api/webhooks/lemonsqueezy"}
    for path, methods in (schema.get("paths") or {}).items():
        if path in PUBLIC_PATHS or not path.startswith("/api/"):
            continue
        for op in methods.values():
            if isinstance(op, dict):
                op.setdefault("security", []).append({"BearerAuth": []})
    app.openapi_schema = schema
    return schema


app.openapi = _custom_openapi


# Mount built frontend last so /api/* routes take precedence. Only mounts when
# the static dir exists — dev mode (Vite on :5173 proxying to us) skips this.
#
# `SPAStaticFiles` falls back to `index.html` on 404 so client-side routes
# (`/documents`, `/projects/:id`, `/account`, `/sign-in/*`, etc) survive a
# direct page load or browser refresh — without it React Router never gets
# to handle those paths because StaticFiles 404s before the fallback.
#
# Modern Starlette raises HTTPException for misses *and* may return a 404
# Response object depending on the path shape (e.g. trailing-slash redirects)
# — handle both. Importing inside the class to keep the symbol scoped.
class _SPAStaticFiles(StaticFiles):
    async def get_response(self, path, scope):
        from starlette.exceptions import HTTPException as _StarletteHTTPException
        try:
            response = await super().get_response(path, scope)
        except _StarletteHTTPException as e:
            if e.status_code == 404:
                return await super().get_response("index.html", scope)
            raise
        if response.status_code == 404:
            return await super().get_response("index.html", scope)
        return response


_static_dir = os.environ.get("STATIC_DIR", "static")
if os.path.isdir(_static_dir):
    app.mount("/", _SPAStaticFiles(directory=_static_dir, html=True), name="static")
    log.info("serving built frontend from %s (SPA fallback enabled)", _static_dir)
