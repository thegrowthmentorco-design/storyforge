"""Export helpers (M6.1).

Currently just DOCX — MD / JSON / CSV are generated client-side because
the frontend already holds the full ExtractionRecord and they're trivial
formats. DOCX needs a library + zipping, so it lives server-side where
python-docx is already installed (also used for upload parsing in
`main._parse_docx`).

Returned as raw bytes so the route layer can wrap in `Response` with the
right `Content-Disposition` header.
"""

from __future__ import annotations

import io
import logging

from docx import Document
from docx.shared import Pt

from db.models import Extraction

log = logging.getLogger("storyforge.exports")


def build_docx(row: Extraction) -> bytes:
    """Assemble a DOCX rendering of an extraction.

    Layout, top to bottom:
      Title (filename)
      Brief — bold "Summary" + paragraph + tag list
      Actors — bullet list
      User stories — H2 per story (id + actor) + As-a/I-want/So-that body
                     + Acceptance criteria bullets + source quote (italic)
      Non-functional requirements — table (Category | Value)
      Gaps & questions — H3 per gap (severity + question) + context paragraph

    Empty sections are skipped (no "Stories: 0 items" placeholder), so a
    document with only a brief still produces a clean export.
    """
    doc = Document()

    # Title
    doc.add_heading(row.filename or "StoryForge extraction", level=0)

    brief = row.brief or {}
    summary = brief.get("summary") or ""
    tags = brief.get("tags") or []

    # Brief
    if summary or tags:
        doc.add_heading("Business summary", level=1)
        if summary:
            doc.add_paragraph(summary)
        if tags:
            tag_p = doc.add_paragraph()
            tag_p.add_run("Tags: ").bold = True
            tag_p.add_run(", ".join(tags))

    # Actors
    actors = row.actors or []
    if actors:
        doc.add_heading("Actors", level=1)
        for a in actors:
            doc.add_paragraph(a, style="List Bullet")

    # Stories
    stories = row.stories or []
    if stories:
        doc.add_heading("User stories", level=1)
        for s in stories:
            sid = s.get("id", "?")
            actor = s.get("actor", "")
            doc.add_heading(f"{sid} — {actor}", level=2)

            body = doc.add_paragraph()
            body.add_run("As a ").bold = True
            body.add_run(actor)
            body.add_run(" I want ").bold = True
            body.add_run(s.get("want", ""))
            body.add_run(" so that ").bold = True
            body.add_run(s.get("so_that", ""))

            section = s.get("section")
            if section:
                p = doc.add_paragraph()
                run = p.add_run(f"Source: {section}")
                run.italic = True
                run.font.size = Pt(9)

            criteria = s.get("criteria") or []
            if criteria:
                p = doc.add_paragraph()
                p.add_run("Acceptance criteria:").bold = True
                for c in criteria:
                    doc.add_paragraph(c, style="List Bullet")

            quote = s.get("source_quote")
            if quote:
                qp = doc.add_paragraph()
                qrun = qp.add_run(f"“{quote}”")
                qrun.italic = True

    # NFRs as a table
    nfrs = row.nfrs or []
    if nfrs:
        doc.add_heading("Non-functional requirements", level=1)
        # 1 header row + N data rows; 2 cols (Category / Value).
        # Style "Light Grid Accent 1" is widely available in Word; falls back
        # to default table styling if missing.
        try:
            table = doc.add_table(rows=1 + len(nfrs), cols=2)
            table.style = "Light Grid Accent 1"
        except Exception:
            table = doc.add_table(rows=1 + len(nfrs), cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Category"
        hdr[1].text = "Value"
        for i, n in enumerate(nfrs, start=1):
            row_cells = table.rows[i].cells
            row_cells[0].text = n.get("category", "")
            row_cells[1].text = n.get("value", "")

    # Gaps
    gaps = row.gaps or []
    if gaps:
        doc.add_heading("Gaps & questions", level=1)
        for g in gaps:
            severity = (g.get("severity") or "?").upper()
            question = g.get("question", "")
            doc.add_heading(f"[{severity}] {question}", level=3)
            context = g.get("context")
            if context:
                doc.add_paragraph(context)
            quote = g.get("source_quote")
            if quote:
                qp = doc.add_paragraph()
                qrun = qp.add_run(f"“{quote}”")
                qrun.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
